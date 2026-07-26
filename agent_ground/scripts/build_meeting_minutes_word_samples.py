from __future__ import annotations

"""회의록 Flow의 TXT 학습 예시를 실제 업로드 가능한 DOCX 예시로 생성한다.

선택한 문서 프리셋은 ``standard_business_brief``이다. 장문의 회의록 원문을
요청한 약 2쪽 안에 유지하기 위해 ``dense_two_page_meeting_minutes``라는
명명된 예외를 적용한다.

- 글꼴: Calibri 대신 한국어 가독성이 좋은 Malgun Gothic
- 여백: 1.0인치 대신 상하 0.34인치, 좌우 0.55인치
- 본문: 11pt/1.10 대신 7.8pt/1.00
- 표 본문: 7.1pt

예외 값은 제목·본문·목록·표·머리글·바닥글에 일관되게 적용하며, 표는
tblW/tblGrid/tcW가 동일한 고정 DXA 구조로 만든다.
"""

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_ROOT = ROOT / "flows" / "meeting_minutes_writer_flow" / "samples"

SOURCE_NAMES = (
    "historical_minutes_01.txt",
    "historical_minutes_02.txt",
)

FONT_NAME = "Malgun Gothic"
INK = "1F2937"
NAVY = "17365D"
BLUE = "2E5F8A"
MUTED = "667085"
LIGHT_FILL = "EEF3F8"
TABLE_BORDER = "B8C5D1"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12_240
LEFT_MARGIN_DXA = 792
RIGHT_MARGIN_DXA = 792
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - LEFT_MARGIN_DXA - RIGHT_MARGIN_DXA
TABLE_INDENT_DXA = 90
CELL_MARGINS_DXA = {"top": 30, "bottom": 30, "start": 90, "end": 90}


def _set_run_font(
    run,
    *,
    size: float,
    bold: bool = False,
    color: str = INK,
    italic: bool = False,
) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float = 1.03,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = line


def _ensure_paragraph_style(
    document: Document,
    name: str,
    *,
    size: float,
    color: str,
    bold: bool = False,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
    keep_with_next: bool = False,
):
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.name = FONT_NAME
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = line
    style.paragraph_format.keep_with_next = keep_with_next
    return style


def _set_style_bottom_border(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "B8C5D1")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell, *, color: str = TABLE_BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def _apply_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("표 열 너비의 합은 본문 너비와 같아야 합니다.")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    _set_width(tbl_pr, "w:tblW", CONTENT_WIDTH_DXA)

    indent = _ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for column_index, width in enumerate(widths_dxa):
        table.columns[column_index].width = Twips(width)

    for row in table.rows:
        row.height = None
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for column_index, cell in enumerate(row.cells):
            width = widths_dxa[column_index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            _set_width(tc_pr, "w:tcW", width)
            cell_margins = _ensure_child(tc_pr, "w:tcMar")
            for side, margin_width in CELL_MARGINS_DXA.items():
                margin = _ensure_child(cell_margins, f"w:{side}")
                margin.set(qn("w:w"), str(margin_width))
                margin.set(qn("w:type"), "dxa")
            _set_cell_border(cell)


def _create_numbering_definition(
    document: Document,
    *,
    format_name: str,
    text: str,
    left_dxa: int,
    hanging_dxa: int,
) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), format_name)
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), text)
    level.append(level_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left_dxa))
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(left_dxa))
    indent.set(qn("w:hanging"), str(hanging_dxa))
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "8")
    spacing.set(qn("w:line"), "242")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_NAME)
    fonts.set(qn("w:hAnsi"), FONT_NAME)
    fonts.set(qn("w:eastAsia"), FONT_NAME)
    run_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), INK)
    run_pr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    run_pr.append(size)
    level.append(run_pr)
    abstract.append(level)
    # OOXML 순서 규칙상 모든 abstractNum은 첫 num 앞에 있어야 한다.
    # 맨 뒤에 넣으면 Word가 사용자 정의 bullet을 기존 decimal처럼 해석할 수 있다.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)


def _clean_inline_text(text: str) -> str:
    return text.replace("`", "").strip()


def _split_markdown(text: str) -> tuple[str, list[str], list[str]]:
    lines = text.replace("\r\n", "\n").splitlines()
    title = ""
    metadata: list[str] = []
    body_start = 0
    for index, raw in enumerate(lines):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            body_start = index + 1
            break
    for index in range(body_start, len(lines)):
        line = lines[index].strip()
        if not line:
            continue
        if line.startswith("## "):
            body_start = index
            break
        if line.startswith("- "):
            metadata.append(line[2:].strip())
    return title, metadata, lines[body_start:]


def _add_title_block(document: Document, title: str, metadata: Iterable[str]) -> None:
    kicker = document.add_paragraph()
    kicker.style = document.styles["AG Kicker"]
    kicker.add_run("AGENT GROUND · 사용자 스타일 학습용 회의록")

    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Title"]
    paragraph.add_run(title)

    metadata_rows = []
    for item in metadata:
        if ":" in item:
            label, value = item.split(":", 1)
            metadata_rows.append((label.strip(), value.strip()))
        else:
            metadata_rows.append(("정보", item))

    table = document.add_table(rows=len(metadata_rows), cols=2)
    widths = [1_120, CONTENT_WIDTH_DXA - 1_120]
    for row_index, (label, value) in enumerate(metadata_rows):
        label_cell, value_cell = table.rows[row_index].cells
        label_cell.text = ""
        value_cell.text = ""
        label_paragraph = label_cell.paragraphs[0]
        value_paragraph = value_cell.paragraphs[0]
        label_paragraph.style = document.styles["AG Metadata Label"]
        value_paragraph.style = document.styles["AG Metadata Value"]
        label_paragraph.add_run(label)
        value_paragraph.add_run(value)
        _set_cell_shading(label_cell, LIGHT_FILL)
        _set_cell_shading(value_cell, WHITE)
    _apply_table_geometry(table, widths)

    spacer = document.add_paragraph()
    _set_paragraph_spacing(spacer, after=0, line=1.0)
    spacer.paragraph_format.space_before = Pt(1.5)


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    if level == 1:
        paragraph.style = document.styles["Heading 1"]
        paragraph.add_run(text)
    else:
        paragraph.style = document.styles["Heading 2"]
        paragraph.add_run(text)


def _add_list_item(document: Document, text: str, *, num_id: int, style_name: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[style_name]
    _apply_numbering(paragraph, num_id)
    paragraph.add_run(_clean_inline_text(text))


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["AG Body"]
    paragraph.add_run(_clean_inline_text(text))


def _is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    body_rows = [row for row in rows[1:] if not _is_separator_row(row)]
    column_count = len(headers)
    table = document.add_table(rows=1 + len(body_rows), cols=column_count)
    all_rows = [headers, *body_rows]

    for row_index, values in enumerate(all_rows):
        for column_index, value in enumerate(values):
            cell = table.rows[row_index].cells[column_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles[
                "AG Table Header" if row_index == 0 else "AG Table Body"
            ]
            paragraph.add_run(_clean_inline_text(value))
            _set_cell_shading(cell, LIGHT_FILL if row_index == 0 else WHITE)

    if column_count == 3:
        widths = [1_100, CONTENT_WIDTH_DXA - 3_100, 2_000]
    elif column_count == 2:
        widths = [2_200, CONTENT_WIDTH_DXA - 2_200]
    else:
        base = CONTENT_WIDTH_DXA // column_count
        widths = [base] * column_count
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    _apply_table_geometry(table, widths)

    trailing = document.add_paragraph()
    _set_paragraph_spacing(trailing, after=0, line=1.0)


def _add_body(document: Document, body_lines: list[str], bullet_num_id: int, decimal_num_id: int) -> None:
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            _add_markdown_table(document, table_rows)
            table_rows = []

    for raw in body_lines:
        line = raw.strip()
        if line.startswith("|") and line.endswith("|"):
            table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
            continue

        flush_table()
        if not line:
            continue
        if line.startswith("## "):
            _add_heading(document, line[3:].strip(), 1)
        elif line.startswith("### "):
            _add_heading(document, line[4:].strip(), 2)
        elif line.startswith("- "):
            _add_list_item(
                document,
                line[2:].strip(),
                num_id=bullet_num_id,
                style_name="AG Bullet",
            )
        elif re.match(r"^\d+\.\s+", line):
            _add_list_item(
                document,
                re.sub(r"^\d+\.\s+", "", line),
                num_id=decimal_num_id,
                style_name="AG Number",
            )
        else:
            _add_body_paragraph(document, line)
    flush_table()


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, fallback, end])
    _set_run_font(run, size=7.5, color=MUTED)


def _configure_document(document: Document, source_name: str) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.34)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(7.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.0

    _ensure_paragraph_style(
        document,
        "AG Kicker",
        size=7.8,
        color=BLUE,
        bold=True,
        after=1,
    )
    _ensure_paragraph_style(
        document,
        "Title",
        size=17.5,
        color=NAVY,
        bold=True,
        after=4,
        keep_with_next=True,
    )
    _ensure_paragraph_style(
        document,
        "AG Metadata Label",
        size=7.5,
        color=BLUE,
        bold=True,
    )
    _ensure_paragraph_style(
        document,
        "AG Metadata Value",
        size=7.5,
        color=INK,
    )
    heading_1_style = _ensure_paragraph_style(
        document,
        "Heading 1",
        size=10.2,
        color=NAVY,
        bold=True,
        before=2.3,
        after=0.8,
        keep_with_next=True,
    )
    _set_style_bottom_border(heading_1_style)
    _ensure_paragraph_style(
        document,
        "Heading 2",
        size=8.5,
        color=BLUE,
        bold=True,
        before=1.7,
        after=0.4,
        keep_with_next=True,
    )
    _ensure_paragraph_style(
        document,
        "AG Body",
        size=7.8,
        color=INK,
        after=0.8,
    )
    _ensure_paragraph_style(document, "AG Bullet", size=7.8, color=INK)
    _ensure_paragraph_style(document, "AG Number", size=7.8, color=INK)
    _ensure_paragraph_style(
        document,
        "AG Table Header",
        size=7.1,
        color=NAVY,
        bold=True,
    )
    _ensure_paragraph_style(document, "AG Table Body", size=7.0, color=INK)

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(header_paragraph, after=0, line=1.0)
    _set_run_font(
        header_paragraph.add_run("사용자 스타일 기반 회의록 작성 Flow · 입력 예시"),
        size=7.2,
        color=MUTED,
    )

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(footer_paragraph, after=0, line=1.0)
    _set_run_font(footer_paragraph.add_run("Agent Ground  |  "), size=7.5, color=MUTED)
    _add_page_number(footer_paragraph)

    document.core_properties.title = source_name.removesuffix(".txt")
    document.core_properties.subject = "Langflow 회의록 작성 Flow용 Word 입력 예시"
    document.core_properties.author = "Agent Ground"
    document.core_properties.comments = (
        "standard_business_brief + dense_two_page_meeting_minutes override; "
        "source text preserved from the paired TXT sample"
    )


def build_word_sample(source_path: Path, output_path: Path) -> None:
    source_text = source_path.read_text(encoding="utf-8")
    title, metadata, body_lines = _split_markdown(source_text)
    if not title or len(metadata) < 2:
        raise ValueError(f"회의록 제목 또는 메타데이터를 찾지 못했습니다: {source_path}")

    document = Document()
    _configure_document(document, source_path.name)
    bullet_num_id = _create_numbering_definition(
        document,
        format_name="bullet",
        text="•",
        left_dxa=300,
        hanging_dxa=180,
    )
    decimal_num_id = _create_numbering_definition(
        document,
        format_name="decimal",
        text="%1.",
        left_dxa=360,
        hanging_dxa=240,
    )
    _add_title_block(document, title, metadata)
    _add_body(document, body_lines, bullet_num_id, decimal_num_id)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    for source_name in SOURCE_NAMES:
        source_path = SAMPLE_ROOT / source_name
        output_path = source_path.with_suffix(".docx")
        build_word_sample(source_path, output_path)
        print(f"CREATED {output_path}")


if __name__ == "__main__":
    main()
